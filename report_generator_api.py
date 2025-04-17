import os
import io
import pandas as pd
import matplotlib.pyplot as plt
import re
import json
import base64
import time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from datetime import datetime
import tempfile
from anthropic import Anthropic

app = Flask(__name__)
CORS(app)  # 啟用跨域請求

# 創建臨時文件夾用於存儲上傳的文件
UPLOAD_FOLDER = tempfile.mkdtemp()
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# 從環境變數讀取 API 金鑰，確保不硬編碼敏感資訊
API_KEY = os.getenv("ANTHROPIC_API_KEY")
if not API_KEY:
    raise ValueError("ANTHROPIC_API_KEY environment variable is not set. Please set it before running the application.")

client = Anthropic(api_key=API_KEY)

# 全局標誌，防止重複下載
is_downloading = False

def get_background_text():
    return (
        "Neonatal inhalation trial: a preliminary feasibility test for drug delivery.\n"
        "The main objective of this experiment is to test whether our company's nebulized drug delivery device can effectively simulate "
        "the clinical neonatal dosage regimen. A549 cells are used as substitutes for primary neonatal cells to evaluate the uniformity of "
        "drug distribution and cellular growth following the nebulization of Survanta and Curosurf into the chip device, providing reference "
        "information for future planning."
    )

def get_files_by_keyword(keyword, valid_exts=(".csv", ".xlsx", ".xls")):
    files_found = {}
    for file in os.listdir(app.config['UPLOAD_FOLDER']):
        if file.lower().endswith(valid_exts) and re.search(keyword.lower(), file.lower()):
            files_found[file] = os.path.join(app.config['UPLOAD_FOLDER'], file)
    return files_found

def read_airflow_files():
    target_files = [
        "0.68Hz_60ml.csv",
        "0.68Hz_500ml.csv",
        "0.75Hz_60ml.csv",
        "0.75Hz_500ml.csv"
    ]
    data = {}
    for file_name in target_files:
        try:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_name)
            if os.path.exists(file_path):
                df = pd.read_csv(file_path)
                if df.empty:
                    print(f"警告：{file_name} 為空數據")
                    continue
                data[file_name] = df
                print(f"讀取 {file_name} 成功，形狀：{df.shape}")
            else:
                # 嘗試查找替代文件
                all_files = [f for f in os.listdir(app.config['UPLOAD_FOLDER']) if f.lower().endswith((".csv", ".xlsx", ".xls"))]
                base_name = file_name.lower().split('.')[0]
                for f in all_files:
                    if base_name in f.lower():
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f)
                        if f.lower().endswith(".csv"):
                            df = pd.read_csv(file_path)
                        else:
                            df = pd.read_excel(file_path)
                        if df.empty:
                            print(f"警告：替代文件 {f} 為空數據")
                            continue
                        data[file_name] = df
                        print(f"找到替代文件 {f} 用於 {file_name}，形狀：{df.shape}")
                        break
        except Exception as e:
            print(f"讀取 {file_name} 錯誤：{e}")

    if not data:
        print("警告：未讀取到任何氣流數據")
    return data

def read_cv_files():
    data = {}
    # 嘗試從 particle_distribution_api.py 獲取匯總 CSV
    particle_api_url = os.getenv("PARTICLE_API_URL", "https://particle-distribution-api.onrender.com")
    particle_analysis_id = request.form.get('particleAnalysisId', '')  # 從前端獲取 particleAnalysisId
    if particle_analysis_id:
        try:
            import requests
            response = requests.get(f"{particle_api_url}/download/{particle_analysis_id}/combined_cv_results.csv")
            if response.status_code == 200:
                df = pd.read_csv(io.BytesIO(response.content))
                if not df.empty:
                    key = f"cv_results_{df['Image'].iloc[0].lower()}"
                    data[key] = df
                    print(f"從 particle_distribution_api 讀取 combined_cv_results.csv 成功，形狀：{df.shape}")
        except Exception as e:
            print(f"從 particle_distribution_api 獲取 CSV 錯誤：{e}")

    # 本地檔案備用
    for file in os.listdir(app.config['UPLOAD_FOLDER']):
        if file.lower().endswith((".csv", ".xlsx", ".xls")) and "cv_results" in file.lower():
            try:
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file)
                if file.lower().endswith(".csv"):
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                if df.empty:
                    print(f"警告：{file} 為空數據")
                    continue
                key = "cv_results_60ml" if "60ml" in file.lower() else "cv_results_500ml"
                data[key] = df
                print(f"讀取 {file} 成功，形狀：{df.shape}")
            except Exception as e:
                print(f"讀取 {file} 錯誤：{e}")
    if not data:
        print("警告：未讀取到任何 CV 數據")
    return data

def read_summary_stats():
    stats_files = get_files_by_keyword("summary_stats")
    if stats_files:
        path = list(stats_files.values())[0]
        try:
            if path.lower().endswith(".csv"):
                df = pd.read_csv(path)
            else:
                df = pd.read_excel(path)
            print(f"讀取 summary_stats 檔案 {path} 成功，形狀：{df.shape}")
            return df
        except Exception as e:
            print(f"讀取 {path} 錯誤：{e}")
            return None
    else:
        print("找不到 summary_stats 檔案。")
        return None

def read_testdata_file():
    test_files = get_files_by_keyword("TestData")
    if test_files:
        path = list(test_files.values())[0]
        try:
            if path.lower().endswith(".csv"):
                df = pd.read_csv(path)
            else:
                df = pd.read_excel(path)
            if df.empty:
                print(f"警告：{path} 為空數據")
                return None
            print(f"讀取 TestData 檔案 {path} 成功，形狀：{df.shape}")
            data = {
                "concentrations": df.iloc[:,0].tolist(),
                "avg_values": df.iloc[:,1].tolist(),
                "std_values": df.iloc[:,2].tolist() if df.shape[1]>=3 else [0]*len(df)
            }
            data["avg_column_name"] = "%Area"
            return {"TestData": data}
        except Exception as e:
            print(f"讀取 {path} 錯誤：{e}")
            return None
    else:
        print("找不到 TestData 檔案。")
        return None

def generate_airflow_wave_chart(file_label, df, subfigure_label):
    try:
        time_col = next((col for col in df.columns if re.search(r'times?[\s\(\)]*(?:ms|毫秒|s|秒)', str(col), re.I)), None)
        flow_col = next((col for col in df.columns if re.search(r'(air)?flow[\s\(\)]*(?:rate)?|氣流', str(col), re.I)), None)
        if time_col is None or flow_col is None:
            print(f"錯誤: {file_label} 無法識別時間或流量欄位")
            return None, None, None

        time_data = pd.to_numeric(df[time_col], errors="coerce").dropna()
        flow_data = pd.to_numeric(df[flow_col], errors="coerce").dropna()

        file_lower = file_label.lower()
        is_068hz = "0.68" in file_lower
        is_500ml = "500" in file_lower
        freq = "0.68Hz" if is_068hz else "0.75Hz"
        vol = "500mL" if is_500ml else "60mL"

        plt.figure(figsize=(4, 3), dpi=150)
        color = 'blue' if is_500ml else 'red'
        plt.plot(time_data, flow_data, color=color, linewidth=1.5)
        plt.grid(False)
        plt.xlim(0, 10)
        plt.xlabel("Time (s)")
        plt.ylabel("Flow Rate (mL/min)")
        plt.tight_layout()

        plt.text(0.05, 0.95, subfigure_label, transform=plt.gca().transAxes, fontsize=12, fontweight='bold')

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', bbox_inches='tight')
        img_stream.seek(0)
        plt.close()

        caption = f"{'Pre' if is_068hz else 'Post'}-treatment ({freq}) {vol} Airflow Pattern"
        return img_stream, caption, subfigure_label
    except Exception as e:
        print(f"產生 {file_label} 氣流波形圖錯誤: {e}")
        return None, None, None

def generate_cv_wave_chart(key_label, df, figure_no):
    try:
        slice_col = next((col for col in df.columns if re.search(r'slice|segment|number', str(col), re.I)), None)
        mean_col = next((col for col in df.columns if re.search(r'mean|intensity', str(col), re.I)), None)
        if slice_col is None or mean_col is None:
            print(f"{key_label} 缺少必要欄位")
            return None, None

        slices = pd.to_numeric(df[slice_col], errors="coerce").dropna()
        mean_vals = pd.to_numeric(df[mean_col], errors="coerce").dropna()

        plt.figure(figsize=(6, 4), dpi=150)
        marker_style = 'o'
        line_color = 'green' if "60ml" in key_label.lower() else 'purple'
        plt.plot(range(len(slices)), mean_vals, marker=marker_style, markersize=5,
                 linestyle='-', linewidth=1.5, color=line_color)
        plt.grid(False)
        plt.xlabel("Flow Channel Position")
        plt.ylabel("Local Mean Intensity")
        x_length = len(slices)
        if x_length > 0:
            positions = [0, x_length//2, x_length-1]
            labels = ["Left", "Middle", "Right"]
            plt.xticks(positions, labels)
            plt.gca().xaxis.set_major_locator(plt.FixedLocator(positions))
        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', bbox_inches='tight')
        img_stream.seek(0)
        plt.close()

        caption = f"Distribution Uniformity for {key_label}"
        return img_stream, caption
    except Exception as e:
        print(f"產生 {key_label} CV 波形圖錯誤: {e}")
        return None, None

def generate_testdata_bar_chart(data, experiment_info, figure_no):
    try:
        plt.figure(figsize=(6, 4), dpi=150)
        plt.grid(False)
        x_labels = data["concentrations"]
        x_idx = list(range(len(x_labels)))
        y_vals = [float(v) for v in data["avg_values"]]
        y_err = [float(v) for v in data["std_values"]]
        plt.bar(x_idx, y_vals, yerr=y_err, capsize=5, alpha=0.8, color='tab:blue',
                ecolor='black', width=0.6)
        plt.xticks(x_idx, x_labels)
        plt.xlabel(f"{experiment_info['compound_name']} Concentration (mg/ml)")
        plt.ylabel(f"{experiment_info['target_protein']} Area (%)")
        plt.ylim(0, 70)
        plt.tight_layout()

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', bbox_inches='tight')
        img_stream.seek(0)
        plt.close()

        caption = f"Concentration-Response of {experiment_info['compound_name']} on {experiment_info['target_protein']}"
        return img_stream, caption
    except Exception as e:
        print(f"生成 TestData 柱狀圖錯誤: {e}")
        return None, None

def generate_airflow_figure_group(doc, airflow_data, figure_no):
    subfigures = []
    subfigure_labels = [f'({chr(97+i)})' for i in range(26)]
    label_idx = 0

    for file_label, df in airflow_data.items():
        if label_idx >= len(subfigure_labels):
            print(f"警告：子圖標記不足，無法為 {file_label} 生成圖表")
            continue
        img_stream, caption, sub_label = generate_airflow_wave_chart(file_label, df, subfigure_labels[label_idx])
        if img_stream:
            subfigures.append((img_stream, caption, sub_label))
            label_idx += 1

    if not subfigures:
        print("無氣流圖表生成")
        return figure_no

    doc.add_paragraph(f"Figure {figure_no}: Airflow Patterns Across Conditions", style='Heading 3')
    rows = (len(subfigures) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.style = 'Table Grid'

    for idx, (img_stream, caption, sub_label) in enumerate(subfigures):
        row = idx // 2
        col = idx % 2
        cell = table.rows[row].cells[col]
        paragraph = cell.add_paragraph()
        paragraph.add_run().add_picture(img_stream, width=Inches(3.0))
        caption_p = cell.add_paragraph(f"Figure {figure_no}{sub_label}: {caption}")
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.runs[0].italic = True

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Figure {figure_no}: Airflow Patterns for Pre- and Post-Treatment Conditions")
    run.bold = True

    return figure_no + 1

def analyze_image_relevance(images, experiment_info, user_prompt):
    if not images:
        return {"groups": [], "standalone": []}

    prompt = f"""
You are analyzing images from a neonatal inhalation trial studying Survanta delivery to A549 cells, targeting SP-B protein. The experiment includes:
- Airflow rate analysis (0.68Hz pre-treatment, 0.75Hz post-treatment, 60mL and 500mL).
- Fluorescent particle distribution uniformity (60mL and 500mL).
- Concentration-response analysis (Survanta doses vs. SP-B fluorescence).

User's Experiment Description:
"{user_prompt}"

Experiment Information:
- Compound: {experiment_info['compound_name']}
- Target Protein: {experiment_info['target_protein']}
- Cell Type: {experiment_info['cell_type']}
- Treatment Time: {experiment_info['treatment_time']}
- Analysis Type: {experiment_info['analysis_type']}

Images provided: {json.dumps(images)}

Task:
- Determine which images are related and should be grouped as a figure group (e.g., multiple microscopy images showing different conditions of the same phenomenon).
- Identify images that should be standalone figures (e.g., unique results like a single overview image).
- For each group or standalone image, suggest a brief caption describing its content.
- Return a JSON object with:
  - "groups": List of groups, each with "images" (list of filenames) and "caption" (string).
  - "standalone": List of standalone images, each with "image" (filename) and "caption" (string).
- Output ONLY valid JSON, with no additional text, comments, or formatting.

Guidelines:
- Images like 'FL_particle.png' likely relate to fluorescent particle distribution.
- Images like 'BF_cell.png' or 'IF_cell_SP-B.png' likely relate to cell morphology or SP-B staining.
- Group images only if they depict closely related conditions (e.g., different doses of the same measurement).
- Assign standalone status to unique or overview images.
- Captions should be concise and professional, suitable for a scientific paper.

Example Output:
{{
  "groups": [
    {{"images": ["FL_particle_60ml.png", "FL_particle_500ml.png"], "caption": "Fluorescent Particle Distribution at Different Volumes"}},
    {{"images": ["BF_cell_1.png", "BF_cell_2.png"], "caption": "Bright-Field Microscopy of A549 Cells"}}
  ],
  "standalone": [
    {{"image": "overview.png", "caption": "Overview of Experimental Setup"}},
    {{"image": "IF_cell_SP-B.png", "caption": "Immunofluorescent Staining of SP-B"}}
  ]
}}
"""

    max_retries = 3
    retry_delay = 5
    for attempt in range(max_retries):
        try:
            response = client.messages.create(
                model="claude-3-7-sonnet-20250219",
                max_tokens=1000,
                temperature=0.3,
                messages=[{"role": "user", "content": prompt}]
            )
            raw_text = response.content[0].text.strip()
            json_start = raw_text.find('{')
            json_end = raw_text.rfind('}') + 1
            if json_start == -1 or json_end == 0:
                raise ValueError("No valid JSON found in response")
            json_text = raw_text[json_start:json_end]
            return json.loads(json_text)
        except Exception as e:
            if "overloaded_error" in str(e) and attempt < max_retries - 1:
                print(f"Claude API 超載，圖片分析嘗試 {attempt + 1}/{max_retries}，等待 {retry_delay} 秒...")
                time.sleep(retry_delay)
                retry_delay *= 2
            else:
                print(f"圖片相關性分析錯誤: {e}")
                groups = []
                standalone = []
                for img in images:
                    img_lower = img.lower()
                    if "fl" in img_lower or "particle" in img_lower:
                        groups.append({
                            "images": [img],
                            "caption": f"Fluorescent Particle Distribution ({img})"
                        })
                    else:
                        standalone.append({
                            "image": img,
                            "caption": f"Cell Image ({img})"
                        })
                return {"groups": groups, "standalone": standalone}

def add_image_to_doc(doc, img_stream, caption, figure_no):
    try:
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_picture(img_stream, width=Inches(6.0))
        caption_p = doc.add_paragraph(f"Figure {figure_no}: {caption}")
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.runs[0].italic = True
        return figure_no + 1
    except Exception as e:
        print(f"添加圖片到報告錯誤: {e}")
        return figure_no

def add_image_group_to_doc(doc, images, caption, figure_no):
    try:
        doc.add_paragraph(f"Figure {figure_no}: {caption}", style='Heading 3')
        rows = (len(images) + 1) // 2
        table = doc.add_table(rows=rows, cols=2)
        table.style = 'Table Grid'

        subfigure_labels = [f'({chr(97+i)})' for i in range(26)]
        for idx, (img_path, sub_caption) in enumerate(images):
            row = idx // 2
            col = idx % 2
            cell = table.rows[row].cells[col]
            paragraph = cell.add_paragraph()
            with open(img_path, 'rb') as f:
                img_stream = io.BytesIO(f.read())
            paragraph.add_run().add_picture(img_stream, width=Inches(3.0))
            caption_p = cell.add_paragraph(f"Figure {figure_no}{subfigure_labels[idx]}: {sub_caption}")
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.runs[0].italic = True

        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_p.add_run(f"Figure {figure_no}: {caption}")
        run.bold = True
        return figure_no + 1
    except Exception as e:
        print(f"添加圖片組到報告錯誤: {e}")
        return figure_no

def analyze_with_claude(concentration_data, experiment_info, user_prompt, airflow_data, cv_data, summary_stats_df):
    background_text = get_background_text()
    data_json = "{}"
    if "TestData" in concentration_data:
        data_json = json.dumps(concentration_data["TestData"], ensure_ascii=False, indent=2)

    airflow_summary = ""
    for f, df in airflow_data.items():
        flow_col = next((c for c in df.columns if re.search(r'flow|rate', str(c), re.I)), None)
        if flow_col:
            avg_flow = pd.to_numeric(df[flow_col], errors='coerce').mean()
            airflow_summary += f"{f} avg flow = {avg_flow:.2f} mL/min\n"

    particle_summary = "N/A"
    if summary_stats_df is not None:
        particle_summary = summary_stats_df.to_string()

    prompt = f"""
I need you to analyze data from a cell experiment and generate a professional academic paper report that equally emphasizes three key components:
1. Airflow rate analysis - comparing neonatal breathing simulation under 0.68Hz (pre-treatment) and 0.75Hz (post-treatment) conditions at both 60mL and 500mL flow.
2. Fluorescent particle distribution uniformity analysis - comparing results from 60mL and 500mL conditions, based on provided CV data and summary statistics.
3. Concentration-response analysis - evaluating the effect of different Survanta doses on SP-B fluorescence intensity from TestData data.

User's Experiment Description:
"{user_prompt}"

Extracted Experimental Information:
- Compound: {experiment_info['compound_name']}
- Target Protein: {experiment_info['target_protein']}
- Cell Type: {experiment_info['cell_type']}
- Treatment Time: {experiment_info['treatment_time']}
- Analysis Type: {experiment_info['analysis_type']}

This study has THREE EQUALLY IMPORTANT components that must be given balanced attention in the report:
    1. Airflow rate analysis - comparing pre-treatment (0.68Hz) and post-treatment (0.75Hz) conditions
    2. Particle distribution uniformity analysis - comparing 500mL and 60mL doses using CV values
    3. Concentration-response analysis - examining the effect of {experiment_info['compound_name']} concentration on {experiment_info['target_protein']} in {experiment_info['cell_type']} cells

Please generate a detailed report in a professional academic paper format, including:

Introduction
- Begin with a comprehensive background on pulmonary surfactant therapy in neonatal applications
- Provide context about {experiment_info['compound_name']} and its relevance in cellular biology
- Briefly explain the importance of all THREE key components: airflow rate simulation, drug distribution uniformity, and concentration-dependent effects
- Include previous research related to these three aspects

Results
- GIVE EQUAL WEIGHT to all three components of the study
- FIRST COMPONENT: Present airflow rate detection results, referring to Figure 1(a)-1(d), comparing pre-treatment and post-treatment conditions, and interpret the patterns observed
- SECOND COMPONENT: Present fluorescent particle distribution results, referring to Figure 2 (60mL) and Figure 3 (500mL), comparing the uniformity between different doses, including data: CV_of_Local_Mean from file {particle_summary}
- THIRD COMPONENT: Present concentration-response analysis results, referring to Figure 4, showing how {experiment_info['compound_name']} affects {experiment_info['target_protein']} distribution
- Do NOT use phrases like "In control cells" - instead, refer to "cells without drug treatment" or "untreated cells"
- Instead of saying percentage values directly (like "167.1%"), express results as "an increase of X%" or "Y fraction of the clinical dose"
- Do NOT use subtitles within the Results section
- Use "mL" as the unit rather than "%" where appropriate
- Provide detailed interpretation of each figure's data, explaining trends and significance

Discussion
- Begin with relevant background information about surfactant therapy
- DISCUSS ALL THREE COMPONENTS with equal emphasis:
  1. The significance of the airflow rate findings for simulating neonatal breathing patterns, referencing Figure 1
  2. The importance of uniform drug distribution for effective therapy, referencing Figure 2 and Figure 3
  3. The concentration-dependent effects on protein distribution, referencing Figure 4
- If additional images are provided (e.g., microscopy images), reference them as Figure 5 onward in the Discussion, describing their relevance to the findings
- Only provide information about drug sources when discussing non-standard applications
- Be conservative in conclusions, basing interpretations strictly on the provided data
- Compare the current findings with previous studies mentioned in the introduction
- Explore possible mechanisms behind the observed effects with caution, avoiding overinterpretation

Conclusion
- Summarize findings from ALL THREE COMPONENTS with balanced emphasis
- Be conservative in your claims, strictly basing conclusions on the observed data
- Highlight potential applications and significance of the findings without overstating implications

Future Directions
- Be conservative in suggesting future directions - avoid mentioning animal experiments, DNA testing, or comparisons between artificial and natural compounds
- Discuss potential future research directions in a balanced, holistic manner
- Suggest broader areas for investigation rather than highly specific experiments
- Consider how improvements in methodology, equipment, or experimental design could enhance understanding
- Propose how this research could be extended to different contexts or applications
- Maintain a balanced approach that doesn't overly focus on any single aspect of the research
- Be conservative in your suggestions, focusing on realistic extensions of the current work

References
- Provide appropriate academic references (at least 3-5) related to the topic, formatted in standard academic citation style
- Include recent papers about all three components: airflow simulation, drug distribution, and {experiment_info['compound_name']} effects
- Use citation markers like [1], [2], etc. in the text

Additional Instructions:
- In the Introduction section, incorporate the following background text (do not include it directly in the prompt):
{background_text}
- In the Results section, begin with airflow measurements (Figure 1), followed by particle deposition (Figure 2, Figure 3), and concentration-response (Figure 4). Use "cells without drug treatment" to refer to untreated cells and report values as relative increases or fractions of the clinical dose, with units in mL.
- In the Discussion section, start by providing background and then compare the results with previous studies. Reference additional images (Figure 5+) if provided, explaining their relevance. Only include drug source details for uncommon usage; conclusions should be conservative and based solely on the provided data.
- In the Future Directions section, propose further in vitro studies only, without suggesting animal experiments, DNA testing, or comparisons between artificial and natural products.

Please write in professional academic English, using formal scientific language with appropriate citations throughout the text. Focus on creating a realistic, high-quality scientific paper that could be published in a respiratory or cell biology journal. ENSURE BALANCED COVERAGE of all three study components and detailed analysis of figures.
Do not include Method in report.
Do not use any Markdown formatting. All section headings must be plain text, on their own line, without any special characters.

You should use these summaries when generating the report:
Airflow Data Summary:
{airflow_summary}

Particle Distribution Summary:
{particle_summary}

Concentration-Response Data (TestData):
{data_json}

Please generate a detailed, balanced academic report in English with the following sections:
Abstract, Introduction, Results, Discussion, Conclusion, Future Directions, and References.
"""

    max_retries = 3
    retry_delay = 5
    for attempt in range(max_retries):
        try:
            response = client.messages.create(
                model="claude-3-7-sonnet-20250219",
                max_tokens=6000,
                temperature=0.3,
                messages=[{"role": "user", "content": prompt}]
            )
            txt = response.content[0].text
            txt = re.sub(r'\*{1,3}', '', txt)
            txt = re.sub(r'\#{1,6}', '', txt)
            return txt
        except Exception as e:
            if "overloaded_error" in str(e) and attempt < max_retries - 1:
                print(f"Claude API 超載，嘗試 {attempt + 1}/{max_retries}，等待 {retry_delay} 秒...")
                time.sleep(retry_delay)
                retry_delay *= 2
            else:
                print(f"Claude API 錯誤: {e}")
                return (
                    "Abstract\nNo content generated due to API error.\n"
                    "Introduction\n...\n"
                    "Results\n...\n"
                    "Discussion\n...\n"
                    "Conclusion\n...\n"
                    "Future Directions\n...\n"
                    "References\n..."
                )

def determine_image_positions(images, experiment_info, user_prompt):
    try:
        relevance = analyze_image_relevance(images, experiment_info, user_prompt)
        positions = {"groups": [], "standalone": []}

        for group in relevance.get("groups", []):
            group_images = []
            for img in group["images"]:
                if img in images and os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], img)):
                    group_images.append((os.path.join(app.config['UPLOAD_FOLDER'], img), img))
                else:
                    print(f"警告：圖片 {img} 不存在，跳過")
            if group_images:
                positions["groups"].append({
                    "images": group_images,
                    "caption": group["caption"]
                })

        for standalone in relevance.get("standalone", []):
            img = standalone["image"]
            if img in images and os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], img)):
                positions["standalone"].append({
                    "image": os.path.join(app.config['UPLOAD_FOLDER'], img),
                    "caption": standalone["caption"]
                })
            else:
                print(f"警告：圖片 {img} 不存在，跳過")

        return positions
    except Exception as e:
        print(f"圖片位置分析錯誤: {e}")
        return {"groups": [], "standalone": []}

def generate_docx_report(claude_report, experiment_info, airflow_data, cv_data, testdata_data, image_positions):
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        doc.add_heading('Experimental Data Analysis Report', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

        report_sections = {}
        current_section = None
        section_content = []
        section_titles = ["Abstract", "Introduction", "Results", "Discussion", "Conclusion", "Future Directions", "References"]
        figure_no = 1

        lines = claude_report.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                if current_section and section_content:
                    section_content.append('')
                continue
            is_section_title = False
            for title in section_titles:
                if line.lower() == title.lower():
                    if current_section:
                        report_sections[current_section] = '\n'.join(section_content)
                    current_section = title
                    section_content = []
                    is_section_title = True
                    break
            if not is_section_title and current_section:
                section_content.append(line)
        if current_section and section_content:
            report_sections[current_section] = '\n'.join(section_content)

        for section_title in section_titles:
            content = report_sections.get(section_title, '')
            if not content.strip():
                continue

            doc.add_paragraph(section_title, style='Heading 1')
            paragraphs = content.split('\n\n')

            if section_title == "Results":
                if airflow_data:
                    doc.add_paragraph("Airflow rate measurements demonstrated the nebulizer's ability to simulate neonatal breathing patterns, as shown in Figure 1.")
                    figure_no = generate_airflow_figure_group(doc, airflow_data, figure_no)

                for key_label, df in cv_data.items():
                    doc.add_paragraph(f"Fluorescent particle distribution uniformity for {key_label} is presented in Figure {figure_no}.")
                    img_stream, caption = generate_cv_wave_chart(key_label, df, figure_no)
                    if img_stream:
                        figure_no = add_image_to_doc(doc, img_stream, caption, figure_no)

                if testdata_data and "TestData" in testdata_data:
                    doc.add_paragraph(f"The concentration-response relationship of {experiment_info['compound_name']} on {experiment_info['target_protein']} is shown in Figure {figure_no}.")
                    img_stream, caption = generate_testdata_bar_chart(testdata_data["TestData"], experiment_info, figure_no)
                    if img_stream:
                        figure_no = add_image_to_doc(doc, img_stream, caption, figure_no)

            elif section_title == "Discussion" and image_positions:
                doc.add_paragraph("Additional imaging data provide further insights into the experimental outcomes, as shown in the following figures.")

                for group in image_positions.get("groups", []):
                    doc.add_paragraph(f"Figure {figure_no} illustrates {group['caption'].lower()}.")
                    figure_no = add_image_group_to_doc(doc, group["images"], group["caption"], figure_no)

                for standalone in image_positions.get("standalone", []):
                    doc.add_paragraph(f"Figure {figure_no} shows {standalone['caption'].lower()}.")
                    with open(standalone["image"], 'rb') as f:
                        img_stream = io.BytesIO(f.read())
                    figure_no = add_image_to_doc(doc, img_stream, standalone["caption"], figure_no)

            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    p = doc.add_paragraph(paragraph_text.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        file_path = os.path.join(app.config['UPLOAD_FOLDER'], "Experimental_Data_Analysis_Report.docx")
        if os.path.exists(file_path):
            os.remove(file_path)  # 刪除舊文件，確保不重複
        doc.save(file_path)
        return file_path
    except Exception as e:
        print(f"生成報告錯誤: {e}")
        return None

# 路由定義 - 健康檢查接口
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "message": "API is running"}), 200

# 上傳文件並分析的接口
@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        # 如果沒有文件被上傳，返回錯誤
        if 'dataFiles' not in request.files:
            return jsonify({"error": "No data files uploaded"}), 400

        # 清空上傳目錄
        for f in os.listdir(app.config['UPLOAD_FOLDER']):
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], f)
            if os.path.isfile(file_path):
                os.remove(file_path)

        # 獲取實驗數據
        experiment_data = {}
        if 'experimentData' in request.form:
            try:
                experiment_data = json.loads(request.form['experimentData'])
            except:
                return jsonify({"error": "Invalid experiment data format"}), 400

        # 設置默認實驗參數
        experiment_info = {
            "compound_name": experiment_data.get('compoundName', 'Survanta'),
            "target_protein": experiment_data.get('targetProtein', 'SP-B'),
            "cell_type": experiment_data.get('cellType', 'A549'),
            "treatment_time": experiment_data.get('treatmentTime', '6小時'),
            "analysis_type": "%Area",
            "signal_unit": "mL"
        }

        # 保存數據文件
        data_files = request.files.getlist('dataFiles')
        saved_data_files = []
        for file in data_files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                saved_data_files.append(filename)

        # 保存圖片文件
        image_files = []
        if 'imageFiles' in request.files:
            files = request.files.getlist('imageFiles')
            for file in files:
                if file and file.filename and file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff')):
                    filename = secure_filename(file.filename)
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    file.save(file_path)
                    image_files.append(filename)

        # 獲取實驗描述
        user_prompt = experiment_data.get('experimentDescription', '')
        if not user_prompt:
            user_prompt = (
                "用霧化器將Survanta霧化加入培養了八天的A549器官晶片中,經過6小時的霧化治療後, 對SP-B進行免疫螢光染色, "
                "TestData - Sheet1.csv是我們在不同Survanta濃度下SP-B螢光分布面積的百分比和標準差, "
                "0.68Hz_60ml.csv, 0.68Hz_500ml.csv,0.75Hz_60ml.csv,0.75Hz_500ml.csv是看霧化裝置是否能模擬臨床給藥的新生兒呼吸頻率, "
                "cv_results_60ml.csv和cv_results_500ml.csv是看此藥方是在器官晶片內，給藥是否均勻，我們用螢光粒子的均勻程度進行評估"
            )

        # 讀取並分析數據
        airflow_data = read_airflow_files()
        cv_data = read_cv_files()
        summary_stats_df = read_summary_stats()
        testdata_data = read_testdata_file()
        concentration_data = testdata_data or {}

        # 分析圖片相關性
        image_positions = None
        if image_files:
            image_positions = determine_image_positions(image_files, experiment_info, user_prompt)

        # 使用 Claude API 生成報告
        claude_report = analyze_with_claude(concentration_data, experiment_info, user_prompt, airflow_data, cv_data, summary_stats_df)

        # 生成 DOCX 報告
        report_path = generate_docx_report(claude_report, experiment_info, airflow_data, cv_data, testdata_data, image_positions)

        if not report_path:
            return jsonify({"error": "Failed to generate report"}), 500

        # 準備報告預覽
        report_preview = claude_report.split("\n\n", 3)[:3]
        report_preview = "\n\n".join(report_preview) + "..."

        # 準備圖表預覽
        chart_previews = []
        # 氣流圖
        if airflow_data:
            for file_name, df in list(airflow_data.items())[:1]:  # 只取第一個作為預覽
                img_stream, caption, _ = generate_airflow_wave_chart(file_name, df, "(a)")
                if img_stream:
                    encoded_img = base64.b64encode(img_stream.getvalue()).decode('utf-8')
                    chart_previews.append({
                        "type": "airflow",
                        "caption": caption,
                        "image": f"data:image/png;base64,{encoded_img}"
                    })

        # 返回結果
        result = {
            "status": "success",
            "message": "報告生成成功",
            "reportPreview": report_preview,
            "chartPreviews": chart_previews,
            "reportPath": os.path.basename(report_path)
        }

        return jsonify(result), 200

    except Exception as e:
        print(f"分析過程中發生錯誤: {e}")
        return jsonify({"error": f"Processing error: {str(e)}"}), 500

# 下載報告的接口
@app.route('/download/<filename>', methods=['GET'])
def download_report(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return jsonify({"error": "File not found"}), 404
    except Exception as e:
        return jsonify({"error": f"Download error: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))